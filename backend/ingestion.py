"""Document ingestion: PDF/DOCX → text chunks with page/section metadata."""
from __future__ import annotations

import io
import re
from pathlib import Path
from typing import BinaryIO

from models import Chunk, Document
from config import settings

# Module-level optional imports — allows patch("ingestion.PdfReader") in tests
try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None  # type: ignore[assignment,misc]

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore[assignment,misc]


def _split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    """Split text into overlapping chunks, respecting sentence boundaries."""
    if not text.strip():
        return []

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Try to break at a sentence boundary (only when not at end of text)
        if end < len(text):
            boundary = max(
                text.rfind(". ", start, end),
                text.rfind(".\n", start, end),
                text.rfind("\n\n", start, end),
            )
            if boundary > start + (chunk_size // 2):
                end = boundary + 1

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        # If we've reached the end of the text, stop
        if end >= len(text):
            break

        # Advance with overlap — but ensure we always move forward
        next_start = end - overlap
        if next_start <= start:
            next_start = end  # prevent stalling or going backward
        start = next_start

    return chunks


def ingest_pdf(file_bytes: bytes, doc_id: str, filename: str) -> tuple[Document, list[Chunk]]:
    """Parse PDF → Document + List[Chunk] with page metadata."""
    if PdfReader is None:
        raise RuntimeError("pypdf not installed. Run: pip install pypdf")

    reader = PdfReader(io.BytesIO(file_bytes))
    chunks: list[Chunk] = []
    chunk_index = 0

    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = re.sub(r"\s+", " ", page_text).strip()
        if not page_text:
            continue

        raw_chunks = _split_text(page_text, settings.chunk_size, settings.chunk_overlap)
        for raw in raw_chunks:
            chunks.append(
                Chunk(
                    doc_id=doc_id,
                    text=raw,
                    page_number=page_num,
                    section_title=f"Page {page_num}",
                    chunk_index=chunk_index,
                )
            )
            chunk_index += 1

    doc = Document(
        doc_id=doc_id,
        filename=filename,
        file_type="pdf",
        total_chunks=len(chunks),
    )
    return doc, chunks


def ingest_docx(file_bytes: bytes, doc_id: str, filename: str) -> tuple[Document, list[Chunk]]:
    """Parse DOCX → Document + List[Chunk] with section/heading metadata."""
    if DocxDocument is None:
        raise RuntimeError("python-docx not installed. Run: pip install python-docx")

    docx = DocxDocument(io.BytesIO(file_bytes))
    chunks: list[Chunk] = []
    chunk_index = 0
    current_section = "Introduction"
    buffer = ""

    def flush_buffer(section: str) -> None:
        nonlocal buffer, chunk_index
        if not buffer.strip():
            return
        raw_chunks = _split_text(buffer.strip(), settings.chunk_size, settings.chunk_overlap)
        for raw in raw_chunks:
            chunks.append(
                Chunk(
                    doc_id=doc_id,
                    text=raw,
                    page_number=None,   # DOCX doesn't expose page nums reliably
                    section_title=section,
                    chunk_index=chunk_index,
                )
            )
            chunk_index += 1
        buffer = ""

    for para in docx.paragraphs:
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")
        text = para.text.strip()
        if not text:
            continue

        if is_heading:
            flush_buffer(current_section)
            current_section = text
        else:
            buffer += " " + text

    flush_buffer(current_section)

    doc = Document(
        doc_id=doc_id,
        filename=filename,
        file_type="docx",
        total_chunks=len(chunks),
    )
    return doc, chunks


def ingest_file(
    file_bytes: bytes,
    filename: str,
    doc_id: str,
) -> tuple[Document, list[Chunk]]:
    """Dispatch to correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return ingest_pdf(file_bytes, doc_id, filename)
    elif ext in (".docx", ".doc"):
        return ingest_docx(file_bytes, doc_id, filename)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")

    reader = PdfReader(io.BytesIO(file_bytes))
    chunks: list[Chunk] = []
    chunk_index = 0

    for page_num, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        page_text = re.sub(r"\s+", " ", page_text).strip()
        if not page_text:
            continue

        raw_chunks = _split_text(page_text, settings.chunk_size, settings.chunk_overlap)
        for raw in raw_chunks:
            chunks.append(
                Chunk(
                    doc_id=doc_id,
                    text=raw,
                    page_number=page_num,
                    section_title=f"Page {page_num}",
                    chunk_index=chunk_index,
                )
            )
            chunk_index += 1

    doc = Document(
        doc_id=doc_id,
        filename=filename,
        file_type="pdf",
        total_chunks=len(chunks),
    )
    return doc, chunks


def ingest_docx(file_bytes: bytes, doc_id: str, filename: str) -> tuple[Document, list[Chunk]]:
    """Parse DOCX → Document + List[Chunk] with section/heading metadata."""
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise RuntimeError("python-docx not installed") from e

    docx = DocxDocument(io.BytesIO(file_bytes))
    chunks: list[Chunk] = []
    chunk_index = 0
    current_section = "Introduction"
    buffer = ""

    def flush_buffer(section: str) -> None:
        nonlocal buffer, chunk_index
        if not buffer.strip():
            return
        raw_chunks = _split_text(buffer.strip(), settings.chunk_size, settings.chunk_overlap)
        for raw in raw_chunks:
            chunks.append(
                Chunk(
                    doc_id=doc_id,
                    text=raw,
                    page_number=None,   # DOCX doesn't expose page nums reliably
                    section_title=section,
                    chunk_index=chunk_index,
                )
            )
            chunk_index += 1
        buffer = ""

    for para in docx.paragraphs:
        style_name = para.style.name if para.style else ""
        is_heading = style_name.startswith("Heading")
        text = para.text.strip()
        if not text:
            continue

        if is_heading:
            flush_buffer(current_section)
            current_section = text
        else:
            buffer += " " + text

    flush_buffer(current_section)

    doc = Document(
        doc_id=doc_id,
        filename=filename,
        file_type="docx",
        total_chunks=len(chunks),
    )
    return doc, chunks


def ingest_file(
    file_bytes: bytes,
    filename: str,
    doc_id: str,
) -> tuple[Document, list[Chunk]]:
    """Dispatch to correct parser based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return ingest_pdf(file_bytes, doc_id, filename)
    elif ext in (".docx", ".doc"):
        return ingest_docx(file_bytes, doc_id, filename)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Only PDF and DOCX are supported.")
